import sys
import os

try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

doc = Document()
doc.add_heading("Classification Report 4", 0)

table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Crop'
hdr_cells[1].text = 'Precision'
hdr_cells[2].text = 'Recall'
hdr_cells[3].text = 'F1-Score'
hdr_cells[4].text = 'Support'

try:
    with open("reports 2/classification_report4.txt", "r") as f:
        lines = f.readlines()
except FileNotFoundError:
    print("reports 2/classification_report4.txt not found.")
    exit(1)

# Adding main classification metrics
for line in lines[4:26]:
    parts = line.split()
    if len(parts) == 5:
        row_cells = table.add_row().cells
        for i in range(5):
            row_cells[i].text = parts[i]

doc.add_paragraph("\n")
doc.add_heading("Overall Metrics", level=2)

table2 = doc.add_table(rows=1, cols=5)
hdr2_cells = table2.rows[0].cells
hdr2_cells[0].text = 'Metric'
hdr2_cells[1].text = 'Precision'
hdr2_cells[2].text = 'Recall'
hdr2_cells[3].text = 'F1-Score'
hdr2_cells[4].text = 'Support'

# Adding accuracy and averages
for line in lines[27:31]:
    if 'accuracy' in line:
        parts = line.split()
        row = table2.add_row().cells
        row[0].text = 'accuracy'
        row[1].text = '-'
        row[2].text = '-'
        row[3].text = parts[1]
        row[4].text = parts[2]
    elif 'avg' in line:
        parts = line.split()
        if len(parts) >= 6:
            row = table2.add_row().cells
            row[0].text = f"{parts[0]} {parts[1]}"
            row[1].text = parts[2]
            row[2].text = parts[3]
            row[3].text = parts[4]
            row[4].text = parts[5]

doc.save("reports 2/classification_report4.docx")
print("✅ Successfully generated Word Document: reports 2/classification_report4.docx")
